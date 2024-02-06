import os
import logging
import traceback
import docx
from flask import Flask, redirect, render_template, request, send_from_directory, make_response, send_file, session, request

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from flask import send_file
from flask_session import Session
from io import BytesIO
from helpers import usd
from datetime import datetime
from jinja2 import Environment, FileSystemLoader

# Configure application
app = Flask(__name__)
app.run(debug=False)


# Configure a log file for errors
handler = logging.FileHandler('error.log')
handler.setLevel(logging.ERROR)
app.logger.addHandler(handler)


# Custom filter
app.jinja_env.filters["usd"] = usd
app.jinja_env.globals["debug"] = True

# Configure session to use filesystem (instead of signed cookies)
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_TYPE"] = "filesystem"
Session(app)



@app.errorhandler(500)
def internal_server_error(error):
    # Log the error details to a file
    app.logger.error("An internal server error occurred", exc_info=True)

    # Create a traceback string for display in the error page
    trace = traceback.format_exc()

    return render_template('500.html', traceback=trace), 500


@app.after_request
def after_request(response):
    """Ensure responses aren't cached"""
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Expires"] = 0
    response.headers["Pragma"] = "no-cache"
    return response


@app.route("/")
def index():
    """Return Index information"""
    return render_template("index.html")



@app.route("/cateringchoices", methods=["GET", "POST"])
def catering_choices():
    """Return Catering Choices information"""
    if request.method == "POST":
        # Handle POST request and return a response
        # You should have some logic here to process the POST data
        # and then return an appropriate response

        return "POST request handled"  # You should replace this with an actual response

    # For GET requests, return a template or render HTML
    return render_template("cateringchoices.html")


# Define functions and classes below
def get_appetizer_price(appetizer_name):
    """Return Appetizer price"""
    # Replace this with your own logic to get the price of the appetizer
    # For example, you could have a dictionary mapping appetizer names to their prices
    appetizer_prices = {
        "Fresh Guacamole w/ House Totopos": 4.49,
        "Chile con Queso w/ House Totopos": 3.49,
        "Hongos al Ajillo": 3.49,
        "Encurtidos": 4.49,
        "Rajas de Chile Poblano": 3.49,
        "Chicken Flautitas": 3.49,
        "Beef Flautitas": 3.49,
        "Bean Flautitas": 3.49,
        "Camarones en Mole Amarillo": 8.99,
        "Chips & Salsa": 2.49,
        "Hand Tossed Salad": 4.49,
    }

    return appetizer_prices.get(appetizer_name, 0.00)


# This function should return the price of the given entree
def get_entree_price(entree_name):
    """Returns the price of the given entree"""
    entree_prices = {
        "Steak Fajitas": 20.99,
        "Chicken Fajitas": 17.99,
        "Shrimp Fajitas": 20.99,
        "Vegetable Fajitas": 17.99,
        "Puntas a la Norteña": 21.99,
        "Mole de Pistachio": 20.99,
        "Mole Coloradito": 20.99,
        "Pork Chile Verde": 14.99,
        "Chicken Chile Verde": 14.99,
        "Chile Colorado": 15.99,
        "Succulent Pork Carnitas": 16.99,
        "Pescado a la Veracruzana": 22.50,
        "Cochinitas Pibil": 16.99,
    }

    return entree_prices.get(entree_name, 0.00)


mole_prices = {
    "Mole Amarillo": 19.99,
    "Mole Poblano": 19.99,
    "Mole Verde": 19.99,
    "Mole Negro": 19.99,
    "Red Pipian": 19.99,
    "Mole Mango": 19.99,
}
a_la_carte_prices = {
    "Traditional Chicken Enchiladas": 5.50,
    "Traditional Beef Enchiladas": 5.50,
    "Traditional Cheese Enchiladas": 5.50,
    "Specialty Enchiladas Suizas": 7.25,
    "Specialty Enchiladas Verde": 7.25,
    "Specialty Enchiladas Mango": 7.25,
    "Specialty Enchiladas Amarillas": 7.25,
    "Enmoladas": 7.25,
    "RI House Chile Rellenos": 7.25,
}
# Define the dessert prices dictionary
dessert_prices = {
    '10" Tres Leches': 32.00,
    '8" Tres Leches': 24.00,
    '10" Flan': 32.00,
    '8" Flan': 24.00,
    '10" Chocolate Flan': 32.00,
    '8" Chocolate Flan': 24.00,
    "Dozen Cookies": 7.50,
}

# Define the drinks prices dictionary
drink_prices = {
    "Assorted Pepsi Products": 2.20,
    "Assorted Jarritos": 2.50,
    "Voss Bottled Water": 3.99,
    "Pints Margarita Mix": 9.99,
    "1/2 Gallons Margarita Mix": 39.96,
    "Gallons Margarita Mix": 79.92,
}

sides_prices = {
    "Spanish Rice & Refried Beans": 3.5,
    "Chips & Salsa": 2.49,
    "Flour Tortillas": 1.0,
    "Corn Tortillas": 1.0,
    "Refried Beans": 1.75,
    "Spanish Rice": 1.75,
    "Black Beans": 1.75,
    "White Rice": 1.75,
}

chafers_prices = {
    "Chafers": 12.0,
    "1/2 Chafers": 6.0,
}


def calculate_total_price(
    appetizers_data,
    entrees_data,
    mole_data,
    a_la_carte_data,
    tamales_data,
    dessert_data,
    drink_data,
    sides_data,
    pints_quarts_data,
    chafers_data,
):
    """Returns the total"""
    appetizer_prices = {
        "Fresh Guacamole w/ House Totopos": 4.49,
        "Chile con Queso w/ House Totopos": 3.49,
        "Hongos al Ajillo": 3.49,
        "Encurtidos": 4.49,
        "Rajas de Chile Poblano": 3.49,
        "Chicken Flautitas": 3.49,
        "Beef Flautitas": 3.49,
        "Bean Flautitas": 3.49,
        "Camarones en Mole Amarillo": 8.99,
        "Chips & Salsa": 2.49,
        "Hand Tossed Salad": 4.49,
    }

    entree_prices = {
        "Steak Fajitas": 20.99,
        "Chicken Fajitas": 17.99,
        "Shrimp Fajitas": 20.99,
        "Vegetable Fajitas": 17.99,
        "Puntas a la Norteña": 21.99,
        "Mole de Pistachio": 20.99,
        "Mole Coloradito": 20.99,
        "Pork Chile Verde": 14.99,
        "Chicken Chile Verde": 14.99,
        "Chile Colorado": 15.99,
        "Succulent Pork Carnitas": 16.99,
        "Pescado a la Veracruzana": 22.50,
        "Cochinitas Pibil": 16.99,
    }

    tamales_prices = {
        "Chicken": {
            "None": 5.5,
            "w/ Chile Verde": 7.25,
            "w/ Mole Amarillo": 7.25,
            "w/ Mole Poblano": 7.25,
            "w/ Mole Verde": 7.25,
            "w/ Mole Negro": 7.25,
            "w/ Red Pipian": 7.25,
            "w/ Mole Mango": 7.25,
            "w/ Mole Coloradito": 7.25,
        },
        "Pork": {
            "None": 5.5,
            "w/ Chile Verde": 7.25,
            "w/ Mole Amarillo": 7.25,
            "w/ Mole Poblano": 7.25,
            "w/ Mole Verde": 7.25,
            "w/ Mole Negro": 7.25,
            "w/ Red Pipian": 7.25,
            "w/ Mole Mango": 7.25,
            "w/ Mole Coloradito": 7.25,
        },
        "Cheese": {
            "None": 5.5,
            "w/ Chile Verde": 7.25,
            "w/ Mole Amarillo": 7.25,
            "w/ Mole Poblano": 7.25,
            "w/ Mole Verde": 7.25,
            "w/ Mole Negro": 7.25,
            "w/ Red Pipian": 7.25,
            "w/ Mole Mango": 7.25,
            "w/ Mole Coloradito": 7.25,
        },
    }

    # Define a dictionary to hold the prices for pints and quarts of mole sauces and salsa selections
    pints_quarts_prices = {
        "Pints": {
            "Salsa": 5.50,
            "Amarillo": 10.99,
            "Poblano": 10.99,
            "Verde": 10.99,
            "Negro": 10.99,
            "Red Pipian": 10.99,
            "Mango": 10.99,
            "Coloradito": 10.99,
        },
        "Quarts": {
            "Salsa": 11.00,
            "Amarillo": 20.99,
            "Poblano": 20.99,
            "Verde": 20.99,
            "Negro": 20.99,
            "Red Pipian": 20.99,
            "Mango": 20.99,
            "Coloradito": 20.99,
        },
        "1/2 Gallon": {
            "Salsa": 22.00,
            "Amarillo": 41.98,
            "Poblano": 41.98,
            "Verde": 41.98,
            "Negro": 41.98,
            "Red Pipian": 41.98,
            "Mango": 41.98,
            "Coloradito": 41.98,
        },
    }

    total_price = 0

    # Calculate the total price for appetizers
    for appetizer, quantity in appetizers_data.items():
        price_per_item = get_appetizer_price(appetizer)
        total_price += price_per_item * quantity

    # Calculate the total price for entrees
    for entree, quantity in entrees_data.items():
        price_per_item = get_entree_price(entree)
        total_price += price_per_item * quantity

    # Calculate the total price for moles
    for mole_flavor, quantity in mole_data.items():
        mole_price = mole_prices[mole_flavor.split(" (")[0]]
        total_price += quantity * mole_price

    for enchilada_name, quantity in a_la_carte_data.items():
        enchilada_price = a_la_carte_prices.get(enchilada_name, 0.00)
        total_price += quantity * enchilada_price

    tamales_total_price = 0  # Initialize tamales_total_price to 0

    for tamales_item in tamales_data:
        protein_choice = tamales_item.get("protein_choice", "none")
        mole_flavor = tamales_item.get(
            "mole_flavor", "none"
        ).strip()  # Clean up mole_flavor
        quantity = tamales_item.get("quantity", 0)
        # Get the index from the tamales_data list
        tamale_index = tamales_data.index(tamales_item)
        # Get the sauce on side option from the form data
        sauce_on_side = bool(request.form.get(f"sauce-on-side[{tamale_index}]"))

        if (
            protein_choice in tamales_prices
            and mole_flavor in tamales_prices[protein_choice]
        ):
            tamales_price = tamales_prices[protein_choice][mole_flavor]
            tamales_total_price += (
                quantity * tamales_price
            )  # Accumulate tamales_total_price
            # Include the sauce on side option in the tamales_data
            tamales_item["sauce_on_side"] = sauce_on_side

    total_price += tamales_total_price  # Add tamales total price to the grand total

    # Calculate Pints & Quarts total price

    pints_quarts_total_price = 0  # Initialize pints_quarts_total_price to 0

        # Inside the Pints & Quarts calculation loop
    for item_data in pints_quarts_data.values():
        pints_quarts_choice = item_data.get("pints_quarts_choice", "")  # Use correct keys here
        pints_quarts_mole_flavor = item_data.get("pints_quarts_mole_flavor", "").strip()  # Use correct keys here
        pints_quarts_quantity = item_data.get("pints_quarts_quantity", 0)  # Use correct keys here

        if pints_quarts_choice in pints_quarts_prices:
            if pints_quarts_mole_flavor in pints_quarts_prices[pints_quarts_choice]:
                pints_quarts_price = (
                    pints_quarts_prices[pints_quarts_choice][pints_quarts_mole_flavor]
                    * pints_quarts_quantity
                )
                pints_quarts_total_price += pints_quarts_price  # Accumulate the total price
            else:
                print(f"DEBUG: mole flavor '{pints_quarts_mole_flavor}' not found in pints_quarts_prices.")
        else:
            print(f"DEBUG: choice '{pints_quarts_choice}' not found in pints_quarts_prices.")
            print(f"DEBUG: pints_quarts_price: {pints_quarts_price}, pints_quarts_quantity: {pints_quarts_quantity}")

    print(f"DEBUG: pints_quarts_total_price: {pints_quarts_total_price}")


    # Move this line outside of the loop to add the total price once after all items are processed.
    total_price += pints_quarts_total_price

    # Calculate dessert total price
    dessert_total_price = sum(
        dessert_prices[dessert] * quantity for dessert, quantity in dessert_data.items()
    )

    drink_total_price = 0

    drink_total_price = sum(
        drink_prices[drink] * quantity for drink, quantity in drink_data.items()
    )
    total_price += drink_total_price

    # Add dessert total price to the grand total
    total_price += dessert_total_price
    sides_total_price = 0
    # Calculate sides total price
    sides_total_price = sum(
        sides_prices[side] * quantity for side, quantity in sides_data.items()
    )
    total_price += sides_total_price

    chafers_total_price = 0
    chafers_total_price = sum(
        chafers_prices[chafer] * quantity for chafer, quantity in chafers_data.items()
    )
    total_price += chafers_total_price
    print(f"DEBUG: Chafers Total Price: {chafers_total_price}")

    return (
        total_price,
        appetizer_prices,
        entree_prices,
        mole_prices,
        dessert_prices,
        drink_prices,
        sides_prices,
        pints_quarts_prices,
        chafers_prices,
    )

# Custom date formatting function
def custom_date_format(order_date):
    day = order_date.day
    if 4 <= day <= 20 or 24 <= day <= 30:
        day_str = str(day) + "th"
    else:
        suffixes = {1: 'st', 2: 'nd', 3: 'rd'}
        day_str = str(day) + suffixes.get(day % 10, 'th')
    
    return order_date.strftime(f"%A, %B {day_str}")


@app.route("/submit_order", methods=["POST"])
def submit_order():
    if request.method == "POST":
        try:
            # Get the contact information from the form
            customer_name = request.form.get("customer_name")
            contact_phone = request.form.get("contact_phone")
            contact_phone = (
                f"{contact_phone[:3]}.{contact_phone[3:6]}.{contact_phone[6:]}"
            )
            print("DEBUG: customer_name:", customer_name)
            print("DEBUG: contact_phone:", contact_phone)

            # Get the customer address from the form
            customer_address = request.form.get("customer_address")
            print("DEBUG: customer_address:", customer_address)

            # Get the date from the form and format it as month day, year
            order_date_str = request.form.get("date")
            order_date = datetime.strptime(order_date_str, "%Y-%m-%d")
            formatted_order_date = custom_date_format(order_date)
            print("DEBUG: formatted_order_date:", formatted_order_date)
            hour = request.form.get("hour")
            minutes = request.form.get("minutes")
            full_time = f"{hour}:{minutes}"
            print("DEBUG: full_time:", full_time)
            serving_for = request.form.get("serving_for")
            pnss_selection = request.form.get("pnss")
            print("DEBUG: serving_for:", serving_for)
            print("DEBUG: pnss_selection:", pnss_selection)

            # Initialize empty data dictionaries
            appetizers_data = {}
            entrees_data = {}
            mole_data = {}
            specialty_chile_rellenos_data = {}
            a_la_carte_data = {}
            tamales_data = []
            a_la_carte_total = 0
            dessert_data = {}
            drink_data = {}
            sides_data = {}
            pints_quarts_data = {}
            chafers_data = {}

            # Iterate over all form items
            for name, quantity in request.form.items():
                if name.startswith("appetizers[") and quantity:
                    appetizer_name = name[len("appetizers[") : -1]
                    appetizers_data[appetizer_name] = int(quantity)
                elif name.startswith("entrees[") and quantity:
                    entree_name = name[len("entrees[") : -1]
                    entrees_data[entree_name] = int(quantity)
                elif name.startswith("mole[") and quantity:
                    mole_name = name[len("mole[") : -1]
                    mole_data[mole_name] = int(quantity)
                elif name.startswith("specialty_chile_rellenos[") and quantity:
                    chile_relleno_name = name[len("specialty_chile_rellenos[") : -1]
                    specialty_chile_rellenos_data[chile_relleno_name] = int(quantity)
                elif name.startswith("a-la-carte[") and quantity:
                    enchilada_name = name[len("a-la-carte[") : -1]
                    a_la_carte_data[enchilada_name] = int(quantity)
                elif name.startswith("drinks[") and quantity:
                    drink_name = name[len("drinks[") : -1]
                    drink_data[drink_name] = int(quantity)  # Store drink data
                elif name.startswith("sides[") and quantity:
                    side_name = name[len("sides[") : -1]
                    sides_data[side_name] = int(quantity)
                    # Handle dessert items
                elif name.startswith("desserts[") and quantity:
                    dessert_name = name[len("desserts[") : -1]
                    dessert_data[dessert_name] = int(quantity)
                elif name.startswith("chafers[") and quantity:
                    chafer_name = name[len("chafers[") : -1]
                    chafers_data[chafer_name] = int(quantity)

                elif name.startswith("tamales-quantity"):
                    tamale_index = int(name.split("[")[1].split("]")[0])
                    print(
                        f"DEBUG: tamale_index: {tamale_index}"
                    )  # Add this line for debugging
                    sauce_on_side = (
                        request.form.get("sauce-on-side[{}]".format(tamale_index))
                        == "on"
                    )

                    if int(quantity) > 0:
                        tamales_data.append(
                            {
                                "protein_choice": request.form.get(
                                    "protein-choice[{}]".format(tamale_index)
                                ),
                                "mole_flavor": request.form.get(
                                    "mole-flavor[{}]".format(tamale_index)
                                ),
                                "quantity": int(quantity),
                                "sauce_on_side": sauce_on_side,
                            }
                        )





            # Iterate through the range of indices (0 to 6 in this case)
            for index in range(7):  # You have 7 items with indices 0 to 6
                # Get the values of the form fields for the current index
                choice = request.form.get(f"pints-quarts-choice[{index}]")
                mole_flavor = request.form.get(f"pints-quarts-mole-flavor[{index}]")
                quantity = request.form.get(f"pints-quarts-quantity[{index}]")

                # Check if all three fields have values (you may want to add additional validation)
                if choice and mole_flavor and quantity:
                    # Add the data to the pints_quarts_data dictionary using the index as the key
                    pints_quarts_data[index] = {
                        "pints_quarts_choice": choice,
                        "pints_quarts_mole_flavor": mole_flavor,
                        "pints_quarts_quantity": int(
                            quantity
                        ),  # Convert quantity to an integer
                    }

           

                    # Retrieve mole data
            mole_flavors = request.form.getlist("mole-flavor[]")
            protein_choices = request.form.getlist("protein-choice[]")
            quantities = request.form.getlist("mole-quantity[]")

            total_price = 0  # Initialize total_price to 0
            for i in range(len(mole_flavors)):
                mole_flavor = mole_flavors[i]
                protein_choice = protein_choices[i]
                quantity = quantities[i]

                if mole_flavor and protein_choice and quantity:
                    mole_data[f"{mole_flavor} ({protein_choice})"] = int(quantity)

            # Calculate a_la_carte_total here after processing a_la_carte_data
            for enchilada, quantity in a_la_carte_data.items():
                price = quantity * a_la_carte_prices[enchilada]
                a_la_carte_total += price
            total_price += a_la_carte_total
            # Iterate over dessert form items
            for name, quantity in request.form.items():
                if name.startswith("desserts[") and quantity:
                    dessert_name = name[len("desserts[") : -1]
                    dessert_data[dessert_name] = int(quantity)

            (
                total_price,
                appetizer_prices,
                entree_prices,
                mole_prices,
                dessert_total,
                drink_prices,
                sides_prices,
                pints_quarts_prices,
                chafers_prices,
            ) = calculate_total_price(
                appetizers_data,
                entrees_data,
                mole_data,
                a_la_carte_data,
                tamales_data,
                dessert_data,
                drink_data,
                sides_data,
                pints_quarts_data,
                chafers_data,
            )

            for chile_relleno_name, quantity in specialty_chile_rellenos_data.items():
                chile_relleno_price = 16.99  # Update the price if it changes
                total_price += quantity * chile_relleno_price

            dessert_total = sum(
                dessert_prices[dessert] * quantity
                for dessert, quantity in dessert_data.items()
            )
            drink_total_price = sum(
                drink_prices[drink] * quantity for drink, quantity in drink_data.items()
            )
            sides_total_price = sum(
                sides_prices[side] * quantity for side, quantity in sides_data.items()
            )
            chafers_total_price = sum(
                chafers_prices[chafer] * quantity
                for chafer, quantity in chafers_data.items()
            )

            # Check if the Full Service checkbox is checked
            full_service_checked = request.form.get("full_service") == "on"
            print(f"DEBUG: full_service_checked: {full_service_checked}")

            # Calculate the Full Service fee based on the checkbox state
            full_service_fee = total_price * 0.37 if full_service_checked else 0
            print(f"DEBUG: full_service_fee: {full_service_fee}")
            print("Form Data:", request.form.to_dict())
            print("DEBUG: full_service_fee (before rendering template):", full_service_fee)

            # Calculate the total price including the Full Service fee
            total_price_with_fee = total_price + full_service_fee
            print(f"DEBUG: total_price_with_fee: {total_price_with_fee}")

            # Define is_tax_exempt based on the checkbox state
            is_tax_exempt = bool(request.form.get("tax_exempt"))
            # Check if the Tax Exempt checkbox is checked
            is_tax_exempt = request.form.get("tax_exempt") == "on"
            # Calculate the packaging and handling fee percentage
            packaging_handling_fee_percentage = request.form.get(
                "packaging_handling_fee", type=float
            )
            print(
                f"DEBUG: packaging_handling_fee_percentage: {packaging_handling_fee_percentage}"
            )

            # Ensure the packaging_handling_fee_percentage is a float and not a string
            if packaging_handling_fee_percentage is not None:
                packaging_handling_fee_percentage = float(
                    packaging_handling_fee_percentage
                )
            else:
                packaging_handling_fee_percentage = (
                    0  # Default to 0 if not provided or invalid
                )
            print(
                f"DEBUG: packaging_handling_fee_percentage (converted): {packaging_handling_fee_percentage}"
            )
            print(
                f"DEBUG: packaging_handling_fee_percentage (after conversion): {packaging_handling_fee_percentage}"
            )

            # Calculate the packaging and handling fee as a percentage of the subtotal
            packaging_handling_fee = total_price * (
                packaging_handling_fee_percentage / 100
            )
            print(f"DEBUG: packaging_handling_fee: {packaging_handling_fee}")

            # Format the packaging and handling fee as a floating-point number with two decimal places
            packaging_handling_fee_formatted = "{:.2f}".format(packaging_handling_fee)

            # Initialize mileage_fee with a default value
            mileage_fee = 0

            # Define the mileage rate
            mileage_rate = 0.67

            # Calculate the mileage fee
            mileage_miles = request.form.get("mileage_fee", type=float)
            print(f"DEBUG: mileage_miles: {mileage_miles}")

            # Ensure mileage_miles is a float and not a string
            if mileage_miles is not None:
                mileage_miles = float(mileage_miles)
                # Calculate the mileage fee as miles multiplied by the rate
                mileage_fee = mileage_miles * mileage_rate * 2
                print(f"DEBUG: mileage_fee: {mileage_fee}")

            print(f"DEBUG: mileage_miles (converted): {mileage_miles}")
            print(f"DEBUG: mileage_miles (after conversion): {mileage_miles}")

            # Calculate the total price with the mileage fee
            total_price_with_mileage_fee = total_price + mileage_fee
            print(
                f"DEBUG: total_price_with_mileage_fee: {total_price_with_mileage_fee}"
            )

            # Calculate the combined total including subtotal, packaging fee, and full service fee
            combined_total = (
                total_price + packaging_handling_fee + full_service_fee + mileage_fee
            )
            print(f"DEBUG: subtotal_plus_packaging_fee: {combined_total}")

            # Calculate tax and total after tax
            tax_rate = 0.0875
            if not is_tax_exempt:
                # Calculate the tax on the combined total
                tax = round(combined_total * tax_rate, 2)
            else:
                tax = 0  # Set tax to 0 if tax exempt
            print(f"DEBUG: tax: {tax}")

            # Calculate the total after tax, including subtotal, packaging fee, full service fee, and tax
            total_after_tax = round(combined_total + tax, 2)
            print(f"DEBUG: total_after_tax: {total_after_tax}")
            print(f"DEBUG: Total Price (Before Tax): {total_price}")
            print(f"DEBUG: Tax: {tax}")
            print(f"DEBUG: Total After Tax: {total_after_tax}")

            # Calculate the gratuity fee percentage
            gratuity_fee_percentage = request.form.get("gratuity", type=float)
            print(f"DEBUG: gratuity_fee_percentage: {gratuity_fee_percentage}")
            # Ensure the gratuity_fee_percentage is a float and not a string
            if gratuity_fee_percentage is not None:
                gratuity_fee_percentage = float(gratuity_fee_percentage)
            else:
                gratuity_fee_percentage = 0  # Default to 0 if not provided or invalid
            print(
                f"DEBUG: gratuity_fee_percentage (converted): {gratuity_fee_percentage}"
            )
            # Calculate the gratuity fee as a percentage of the total after tax
            gratuity_fee = (
                total_after_tax * (gratuity_fee_percentage / 100)
                if gratuity_fee_percentage is not None
                else 0
            )
            print(f"DEBUG: gratuity_fee: {gratuity_fee}")
            # Format the gratuity fee as a floating-point number with two decimal places
            gratuity_fee_formatted = "{:.2f}".format(gratuity_fee)

            # Handle chafers data
            total_chafers = request.form.get("total_chafers")
            chafers_quantity = request.form.get("chafers[Chafers]", type=int) or 0
            half_chafers_quantity = (
                request.form.get("chafers[1/2 Chafers]", type=int) or 0
            )
            # Calculate the total quantity of chafers
            # Calculate total_chafers as a boolean value
            total_chafers = any(
                [int(chafers_quantity), int(half_chafers_quantity)]
            )  # Check if any chafers are selected
            print(f"DEBUG: total_chafers: {total_chafers}")
            # Add this line to your app.py just before rendering the template
            print(
                f"DEBUG: full_service_fee before rendering template: {full_service_fee}"
                    )
       


        # Store the entire context in the Flask session
            session['order_data'] = {
                'customer_name': customer_name,
                'formatted_order_date': formatted_order_date,
                'contact_phone': contact_phone,
                'customer_address': customer_address,
                'serving_for': serving_for,
                'full_time': full_time,
                'appetizers_data': appetizers_data,
                'entrees_data': entrees_data,
                'mole_data': mole_data,
                'specialty_chile_rellenos_data': specialty_chile_rellenos_data,
                'a_la_carte_data': a_la_carte_data,
                'tamales_data': tamales_data,
                'dessert_data': dessert_data,
                'drink_data': drink_data,
                'sides_data': sides_data,
                'pints_quarts_data': pints_quarts_data,
                'chafers_data': chafers_data,
                'appetizer_prices': appetizer_prices,
                'entree_prices': entree_prices,
                'mole_prices': mole_prices,
                'a_la_carte_total': a_la_carte_total,
                'a_la_carte_prices': a_la_carte_prices,
                'dessert_total': dessert_total,
                'dessert_prices': dessert_prices,
                'drink_total_price': drink_total_price,
                'drink_prices': drink_prices,
                'sides_total_price': sides_total_price,
                'sides_prices': sides_prices,
                'pints_quarts_prices': pints_quarts_prices,
                'chafers_total_price': chafers_total_price,
                'chafers_prices': chafers_prices,
                'order_date': formatted_order_date,
                'subtotal': total_price,
                'tax': tax,
                'total_after_tax': total_after_tax,
                'delivery_selection': request.form.get("delivery_selection"),
                'pickup_location_selection': request.form.get("pickup_location_selection"),
                'total_chafers': total_chafers,
                'full_service_fee': full_service_fee,
                'total_price_with_fee': total_price_with_fee,
                'is_tax_exempt': is_tax_exempt,
                'packaging_handling_fee': packaging_handling_fee,
                'packaging_handling_fee_percentage': packaging_handling_fee_percentage,
                'combined_total': combined_total,
                'packaging_handling_fee_formatted': packaging_handling_fee_formatted,
                'gratuity_fee_formatted': gratuity_fee_formatted,
                'gratuity_fee_percentage': gratuity_fee_percentage,
                'gratuity_fee': gratuity_fee,
                'mileage_fee': mileage_fee,
                'total_price_with_mileage_fee': total_price_with_mileage_fee,
                'mileage_rate': mileage_rate,
                'pnss_selection': pnss_selection
            }




            return render_template(
                "success.html",
                customer_name=customer_name,
                contact_phone=contact_phone,
                customer_address=customer_address,
                serving_for=serving_for,
                full_time=full_time,
                appetizers_data=appetizers_data,
                entrees_data=entrees_data,
                mole_data=mole_data,
                specialty_chile_rellenos_data=specialty_chile_rellenos_data,
                a_la_carte_data=a_la_carte_data,
                tamales_data=tamales_data,
                dessert_data=dessert_data,
                drink_data=drink_data,
                sides_data=sides_data,
                pints_quarts_data=pints_quarts_data,
                chafers_data=chafers_data,
                appetizer_prices=appetizer_prices,
                entree_prices=entree_prices,
                mole_prices=mole_prices,
                a_la_carte_total=a_la_carte_total,
                a_la_carte_prices=a_la_carte_prices,
                dessert_total=dessert_total,
                dessert_prices=dessert_prices,
                drink_total_price=drink_total_price,
                drink_prices=drink_prices,
                sides_total_price=sides_total_price,
                sides_prices=sides_prices,
                pints_quarts_prices=pints_quarts_prices,
                chafers_total_price=chafers_total_price,
                chafers_prices=chafers_prices,
                order_date=formatted_order_date,
                subtotal=total_price,
                tax=tax,
                total_after_tax=total_after_tax,
                delivery_selection=request.form.get("delivery_selection"),
                pickup_location_selection=request.form.get("pickup_location_selection"),
                total_chafers=total_chafers,
                full_service_fee=full_service_fee,
                total_price_with_fee=total_price_with_fee,
                is_tax_exempt=is_tax_exempt,
                packaging_handling_fee=packaging_handling_fee,
                packaging_handling_fee_percentage=packaging_handling_fee_percentage,
                combined_total=combined_total,
                packaging_handling_fee_formatted=packaging_handling_fee_formatted,
                gratuity_fee_formatted=gratuity_fee_formatted,
                gratuity_fee_percentage=gratuity_fee_percentage,
                gratuity_fee=gratuity_fee,
                mileage_fee=mileage_fee,
                total_price_with_mileage_fee=total_price_with_mileage_fee,
                mileage_rate=mileage_rate,
                pnss_selection=pnss_selection
                
            )

        except Exception as e:
            print("ERROR:", e)
            return "An error occurred while processing the order.", 500

    # If it's a GET request, redirect to the form page (optional)
    return redirect("/order_form")







@app.route('/download_success_docx')
def download_success_docx():
    # Create a new Word document
    doc = Document()

    # Adjust the top margin for the entire document (section properties)
    section = doc.sections[0]  # Access the first section (there's usually only one)
    section.top_margin = Inches(0.25)  # Set the top margin to 0.25 inches (adjust as needed)

    # Retrieve the form data from session variables
    order_data = session.get('order_data', {})
    customer_name = order_data.get('customer_name')
    formatted_order_date = order_data.get('formatted_order_date')
    contact_phone = order_data.get('contact_phone')
    customer_address = order_data.get('customer_address')
    serving_for = order_data.get('serving_for')
    full_time = order_data.get('full_time')
    appetizers_data = order_data.get('appetizers_data')
    entrees_data = order_data.get('entrees_data')
    mole_data = order_data.get('mole_data')
    specialty_chile_rellenos_data = order_data.get('specialty_chile_rellenos_data')
    a_la_carte_data = order_data.get('a_la_carte_data')
    tamales_data = order_data.get('tamales_data')
    dessert_data = order_data.get('dessert_data')
    drink_data = order_data.get('drink_data')
    sides_data = order_data.get('sides_data')
    pints_quarts_data = order_data.get('pints_quarts_data')
    chafers_data = order_data.get('chafers_data')
    appetizer_prices = order_data.get('appetizer_prices')
    entree_prices = order_data.get('entree_prices')
    mole_prices = order_data.get('mole_prices')
    a_la_carte_total = order_data.get('a_la_carte_total')
    a_la_carte_prices = order_data.get('a_la_carte_prices')
    dessert_total = order_data.get('dessert_total')
    dessert_prices = order_data.get('dessert_prices')
    drink_total_price = order_data.get('drink_total_price')
    drink_prices = order_data.get('drink_prices')
    sides_total_price = order_data.get('sides_total_price')
    sides_prices = order_data.get('sides_prices')
    pints_quarts_prices = order_data.get('pints_quarts_prices')
    chafers_total_price = order_data.get('chafers_total_price')
    chafers_prices = order_data.get('chafers_prices')
    order_date = order_data.get('order_date')
    subtotal = order_data.get('subtotal')
    tax = order_data.get('tax')
    total_after_tax = order_data.get('total_after_tax')
    delivery_selection = order_data.get('delivery_selection')
    pickup_location_selection = order_data.get('pickup_location_selection')
    total_chafers = order_data.get('total_chafers')
    full_service_fee = order_data.get('full_service_fee')
    total_price_with_fee = order_data.get('total_price_with_fee')
    is_tax_exempt = order_data.get('is_tax_exempt')
    packaging_handling_fee = order_data.get('packaging_handling_fee')
    packaging_handling_fee_percentage = order_data.get('packaging_handling_fee_percentage')
    combined_total = order_data.get('combined_total')
    packaging_handling_fee_formatted = order_data.get('packaging_handling_fee_formatted')
    gratuity_fee_formatted = order_data.get('gratuity_fee_formatted')
    gratuity_fee_percentage = order_data.get('gratuity_fee_percentage')
    gratuity_fee = order_data.get('gratuity_fee')
    mileage_fee = order_data.get('mileage_fee')
    total_price_with_mileage_fee = order_data.get('total_price_with_mileage_fee')
    mileage_rate = order_data.get('mileage_rate')
    pnss_selection = order_data.get('pnss_selection')

    
    # Define a common style for text
    common_style = doc.styles['Normal']
    font = common_style.font
    font.name = 'Arial'
    font.size = Pt(11)

        # Set margins and spacing for the 'Normal' style
    common_style.paragraph_format.line_spacing = 1.15
    common_style.paragraph_format.space_before = Pt(0.5)
    common_style.paragraph_format.space_after = Pt(0)
    common_style.paragraph_format.left_indent = Pt(0.5)
    common_style.paragraph_format.right_indent = Pt(0.5)

    # Add the customer name with custom formatting
    customer_info_paragraph = doc.add_paragraph()
    customer_info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    customer_info_run = customer_info_paragraph.add_run(f'{customer_name}')
    customer_info_run.bold = True


            # Add the "Delivery Selection" if selected
    if delivery_selection:
        print("Delivery:", delivery_selection)
        customer_info_paragraph.add_run(' ').bold = True  # Add a space for separation
        customer_info_paragraph.add_run(f'{delivery_selection}').bold = True
            # Add any specific formatting for the delivery selection here

        # Add the "Pickup Location" if selected
    if pickup_location_selection:
        print("Pickup:", pickup_location_selection)
        customer_info_paragraph.add_run(' ')  # Add space for separation
        customer_info_paragraph.add_run(f'{pickup_location_selection}').bold = True

    # Add the "Full Service" if selected
    if full_service_fee:
        print("Full Service: Yes")
        customer_info_paragraph.add_run(' Full Service').bold = True
        
    # Add the "w/ Chafers" if selected
    if total_chafers:
        print("w/ Chafers")
        customer_info_paragraph.add_run(' w/ Chafers').bold = True

    # Add the Date and Time of the event
    formatted_order_date_paragraph = doc.add_paragraph()
    formatted_order_date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    formatted_order_date_run = formatted_order_date_paragraph.add_run(f'{formatted_order_date}   Arrival: {full_time}')
    formatted_order_date_run.bold = True


    # Add the customer address with custom formatting
    customer_address_paragraph = doc.add_paragraph()
    customer_address_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    customer_address_run = customer_address_paragraph.add_run(f'Address: {customer_address}')
    customer_address_run.bold = True

    # Add the contact phone with custom formatting
    contact_phone_paragraph = doc.add_paragraph()
    contact_phone_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_phone_run = contact_phone_paragraph.add_run(f'Contact #: {contact_phone}')
    contact_phone_run.bold = True

    # Add a blank line
    doc.add_paragraph("", style='Normal')

    # Add How many people are being served
    serving_for_paragraph = doc.add_paragraph()
    serving_for_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    serving_for_run = serving_for_paragraph.add_run(f'Serving {serving_for}')
    serving_for_run.bold = True

    # Add the PNSS information
    pnss_paragraph = doc.add_paragraph()
    pnss_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if pnss_selection:
        print("No PNSS")
        pnss_run = pnss_paragraph.add_run('No PNSS')
        pnss_run.font.size = Pt(11)
        pnss_run.font.italic = True
        pnss_run.font.color.rgb = RGBColor(141, 0, 0)
        pnss_paragraph.paragraph_format.space_after = Pt(1)
        pnss_run.bold = True
    else:
        print("Providing quality disposable plates, napkins, silverware, serving utensils")
        pnss_run = pnss_paragraph.add_run('Providing quality disposable plates, napkins, silverware, serving utensils')
        pnss_run.font.size = Pt(11)
        pnss_run.font.italic = True
        pnss_run.font.color.rgb = RGBColor(141, 0, 0)
        pnss_run.bold = True






    
    # Add a section for Appetizers
    if appetizers_data:
        # Add a blank line
        doc.add_paragraph("", style='Normal')

        # Add the "Appetizers" section header
        appetizers_paragraph = doc.add_paragraph()
        appetizers_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        appetizers_run = appetizers_paragraph.add_run('~Appetizers~')
        appetizers_run.bold = True
        appetizers_run.font.size = Pt(14)

        # Variables to track Flautitas and Salad
        flautitas_found = False
        flautitas_data = []
        salad_data = []

        # Add appetizers details
        for appetizer, quantity in appetizers_data.items():
            if 'Flautitas' in appetizer:
                flautitas_found = True
                flautitas_name = appetizer.split(' for')[0]
                flautitas_data.append({'name': flautitas_name, 'quantity': quantity})
            elif appetizer == 'Hand Tossed Salad':
                salad_data.append({'name': appetizer, 'quantity': quantity})
            else:
                # For other appetizers, print the name and quantity
                appetizer_ordered = f'{appetizer} for {quantity}... {"%.2f" % (quantity * appetizer_prices[appetizer])}'
                appetizer_paragraph = doc.add_paragraph(appetizer_ordered, style='Normal')
                appetizer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Check if Flautitas were found
        if flautitas_found:
            # Add Flautitas details
            for flautitas in flautitas_data:
                flautitas_ordered = f'{flautitas["name"]} for {flautitas["quantity"]}... {"%.2f" % (flautitas["quantity"] * appetizer_prices[flautitas["name"]])}'
                flautitas_paragraph = doc.add_paragraph(flautitas_ordered, style='Normal')
                flautitas_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add a line under the last Flautitas selection
            crema_line = doc.add_paragraph('w/ garnishes of crema mexicana, salsa de aguacate, and queso cotija', style='Normal')
            crema_line.runs[0].font.size = Pt(11)
            crema_line.runs[0].font.italic = True
            crema_line.runs[0].font.color.rgb = RGBColor(141, 0, 0)
            crema_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Bold the crema line
            for run in crema_line.runs:
                run.bold = True

        # Add Salad details
        for salad in salad_data:
            salad_ordered = f'{salad["name"]} for {salad["quantity"]}... {"%.2f" % (salad["quantity"] * appetizer_prices[salad["name"]])}'
            salad_paragraph = doc.add_paragraph(salad_ordered, style='Normal')
            salad_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Print the salad line directly under Hand Tossed Salad
            salad_line = doc.add_paragraph('w/ dressing and cheese on the side', style='Normal')
            salad_line.runs[0].font.size = Pt(11)
            salad_line.runs[0].font.italic = True
            salad_line.runs[0].font.color.rgb = RGBColor(141, 0, 0)
            salad_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Bold the salad line
            for run in salad_line.runs:
                run.bold = True





    # Add a section for Entrees, Moles, and Specialty Chile Rellenos
    if entrees_data or mole_data or specialty_chile_rellenos_data:
        # Add a blank line
        doc.add_paragraph("", style='Normal')

        # Add the "Entrees" section header
        entrees_paragraph = doc.add_paragraph()
        entrees_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        entrees_run = entrees_paragraph.add_run('~Entrees~')
        entrees_run.bold = True
        entrees_run.font.size = Pt(14)

        # Add description for Entrees
        entrees_description = doc.add_paragraph("Entrees include chips, salsa, Spanish rice, refried beans, and tortillas.",
                                            style='Normal')
        entrees_description.runs[0].font.size = Pt(11)
        entrees_description.runs[0].font.italic = True
        entrees_description.runs[0].font.color.rgb = RGBColor(141, 0, 0)
        entrees_description.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        entrees_description.runs[0].bold = True  # Make the description bold
        

        # Variables to track Fajitas, Cochinitas Pibil, and Carnitas
    fajitas_found = False
    fajitas_data = []
    cochinitas_found = False
    cochinitas_data = []
    carnitas_found = False
    carnitas_data = []

    # Add entrees details
    for entree, quantity in entrees_data.items():
        if 'Fajitas' in entree:
            fajitas_found = True
            fajitas_name = entree.split(' for')[0]
            fajitas_data.append({'name': fajitas_name, 'quantity': quantity})
        elif entree == 'Cochinitas Pibil':
            cochinitas_found = True
            cochinitas_data.append({'name': entree, 'quantity': quantity})
        elif entree == 'Succulent Pork Carnitas':
            carnitas_found = True
            carnitas_data.append({'name': entree, 'quantity': quantity})
        else:
            # For other entrees, print the name and quantity
            entree_ordered = f'{entree} for {quantity}... {"%.2f" % (quantity * entree_prices[entree])}'
            entree_paragraph = doc.add_paragraph(entree_ordered, style='Normal')
            entree_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Check if Fajitas were found
    if fajitas_found:
        # Add Fajitas details
        for fajitas in fajitas_data:
            fajitas_ordered = f'{fajitas["name"]} for {fajitas["quantity"]}... {"%.2f" % (fajitas["quantity"] * entree_prices[fajitas["name"]])}'
            fajitas_paragraph = doc.add_paragraph(fajitas_ordered, style='Normal')
            fajitas_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a line under the last Fajitas selection
        fajitas_line = doc.add_paragraph('Fajitas accompanied by guacamole, pico de gallo, sour cream, jalapeños, and queso cotija', style='Normal')
        fajitas_line.runs[0].font.size = Pt(10)
        fajitas_line.runs[0].font.italic = True
        fajitas_line.runs[0].font.color.rgb = RGBColor(141, 0, 0)
        fajitas_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Bold the Fajitas line
        for run in fajitas_line.runs:
            run.bold = True

    # Check if Cochinitas Pibil were found
    if cochinitas_found:
        # Add Cochinitas Pibil details
        for cochinitas in cochinitas_data:
            cochinitas_ordered = f'{cochinitas["name"]} for {cochinitas["quantity"]}... {"%.2f" % (cochinitas["quantity"] * entree_prices[cochinitas["name"]])}'
            cochinitas_paragraph = doc.add_paragraph(cochinitas_ordered, style='Normal')
            cochinitas_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a line under the last Cochinitas Pibil selection
        cochinitas_line = doc.add_paragraph('w/ pickled red onions', style='Normal')
        cochinitas_line.runs[0].font.size = Pt(11)
        cochinitas_line.runs[0].font.italic = True
        cochinitas_line.runs[0].font.color.rgb = RGBColor(141, 0, 0)
        cochinitas_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Bold the Cochinitas Pibil line
        for run in cochinitas_line.runs:
            run.bold = True

    # Check if Succulent Pork Carnitas were found
    if carnitas_found:
        # Add Succulent Pork Carnitas details
        for carnitas in carnitas_data:
            carnitas_ordered = f'{carnitas["name"]} for {carnitas["quantity"]}... {"%.2f" % (carnitas["quantity"] * entree_prices[carnitas["name"]])}'
            carnitas_paragraph = doc.add_paragraph(carnitas_ordered, style='Normal')
            carnitas_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a line under the last Succulent Pork Carnitas selection
        carnitas_line = doc.add_paragraph('Accompanied by garnishes of pico de gallo, salsa de aguacate, sliced jalapeños', style='Normal')
        carnitas_line.runs[0].font.size = Pt(11)
        carnitas_line.runs[0].font.italic = True
        carnitas_line.runs[0].font.color.rgb = RGBColor(141, 0, 0)
        carnitas_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Bold the Carnitas line
        for run in carnitas_line.runs:
            run.bold = True


                    # Add Mole details
        for mole_flavor, quantity in mole_data.items():
            mole_info = mole_flavor.split(' (', 1)
            mole_name = mole_info[0]
            protein_choice = mole_info[1][:-1] if len(mole_info) > 1 else ''
            price = quantity * mole_prices[mole_name.strip()]
            mole_ordered = f'{mole_name} {protein_choice} for {quantity}... {"{:.2f}".format(price)}'
            mole_paragraph = doc.add_paragraph(mole_ordered, style='Normal')
            mole_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Iterate through the specialty_chile_rellenos_data dictionary to display specialty chiles
        for chile, quantity in specialty_chile_rellenos_data.items():
            chile_ordered = f'{chile} for {quantity}... {"{:.2f}".format(quantity * 16.99)}'
            chile_paragraph = doc.add_paragraph(chile_ordered, style='Normal')
            chile_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER





    # Add a section for A la Carte
    if a_la_carte_data or tamales_data:
        # Add a blank line
        doc.add_paragraph("", style='Normal')

        # Add the "A la Carte" section header
        a_la_carte_paragraph = doc.add_paragraph()
        a_la_carte_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        a_la_carte_run = a_la_carte_paragraph.add_run('~A la Carte~')
        a_la_carte_run.bold = True
        a_la_carte_run.font.size = Pt(14)

        # Variables to track A la Carte total
        a_la_carte_total = 0

        # Add A la Carte details
        for enchilada, quantity in a_la_carte_data.items():
            price = quantity * a_la_carte_prices[enchilada]
            enchilada_ordered = f'{quantity} {enchilada}... {"{:.2f}".format(price)}'
            enchilada_paragraph = doc.add_paragraph(enchilada_ordered, style='Normal')
            enchilada_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            a_la_carte_total += price

        # Iterate through Tamales data
        for tamales_item in tamales_data:
            if tamales_item['quantity'] > 0:
                tamales_price = 7.25 if tamales_item['mole_flavor'] != 'None' else 5.50
                tamales_item_total = tamales_price * tamales_item['quantity']
                tamales_ordered = (
                    f'{tamales_item["quantity"]} {tamales_item["protein_choice"]} Tamales'
                    f'{f" with {tamales_item["mole_flavor"]}" if tamales_item["mole_flavor"] != 'None' else ''}'
                    f'{f" on side" if tamales_item["sauce_on_side"] else ""}'
                    f'... {"{:.2f}".format(tamales_item_total)}'
                )
                tamales_paragraph = doc.add_paragraph(tamales_ordered, style='Normal')
                tamales_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                a_la_carte_total += tamales_item_total


        # Check if there is dessert data
    if dessert_data:
        # Add a blank line
        doc.add_paragraph("", style='Normal')

        # Add the "Desserts" section header
        dessert_paragraph = doc.add_paragraph()
        dessert_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        dessert_run = dessert_paragraph.add_run('~Desserts~')
        dessert_run.bold = True
        dessert_run.font.size = Pt(14)

        # Create a list to store dessert details
        dessert_details = []

        # Add dessert details
        for dessert, quantity in dessert_data.items():
            price = quantity * dessert_prices[dessert]
            if dessert == 'Dozen Cookies':
                dessert_details.append(f'{quantity} {dessert}... {"{:.2f}".format(price)}')
            else:
                dessert_details.append(f'{quantity} ea. {dessert}... {"{:.2f}".format(price)}')

        # Add dessert details to the document
        for detail in dessert_details:
            dessert_paragraph = doc.add_paragraph(detail, style='Normal')
            dessert_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


            # Check if there is drink data
    if drink_data:
        # Add a blank line
        doc.add_paragraph("", style='Normal')

        # Add the "Drinks" section header
        drink_paragraph = doc.add_paragraph()
        drink_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        drink_run = drink_paragraph.add_run('~Drinks~')
        drink_run.bold = True
        drink_run.font.size = Pt(14)

        # Add drink details
        for drink, quantity in drink_data.items():
            price = quantity * drink_prices[drink]
            drink_ordered = f'{quantity} {drink}... {"{:.2f}".format(price)}'
            drink_paragraph = doc.add_paragraph(drink_ordered, style='Normal')
            drink_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Check for special instructions for Margarita Mix
            if quantity > 1 and drink in ['Pints Margarita Mix', '1/2 Gallons Margarita Mix', 'Gallons Margarita Mix']:
                special_instructions = f'w/ salt and lime wedges'
            elif drink in ['Pints Margarita Mix', '1/2 Gallons Margarita Mix', 'Gallons Margarita Mix']:
                special_instructions = 'w/ salt and lime wedges'
            else:
                special_instructions = None

            # Add special instructions if applicable
            if special_instructions:
                special_instructions_paragraph = doc.add_paragraph(special_instructions, style='Normal')
                special_instructions_paragraph.runs[0].font.size = Pt(11)
                special_instructions_paragraph.runs[0].font.italic = True
                special_instructions_paragraph.runs[0].font.color.rgb = RGBColor(128, 0, 0)
                special_instructions_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    # Check if there is sides data
    if sides_data:
        # Add a blank line
        doc.add_paragraph("", style='Normal')

        # Add the "Sides" section header
        sides_paragraph = doc.add_paragraph()
        sides_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sides_run = sides_paragraph.add_run('~Sides~')
        sides_run.bold = True
        sides_run.font.size = Pt(14)

        # Add side details
        for side, quantity in sides_data.items():
            price = quantity * sides_prices[side]
            side_ordered = f'{side} for {quantity}... {"{:.2f}".format(price)}'
            side_paragraph = doc.add_paragraph(side_ordered, style='Normal')
            side_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    # Check if there is pints_quarts or chafers data
    if pints_quarts_data or chafers_data:
        # Add a blank line
        doc.add_paragraph("", style='Normal')

        # Add the "Extras" section header
        extras_paragraph = doc.add_paragraph()
        extras_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        extras_run = extras_paragraph.add_run('~Extras~')
        extras_run.bold = True
        extras_run.font.size = Pt(14)

        # Add pints_quarts details
    if pints_quarts_data:
        for item_id, item_data in pints_quarts_data.items():
            # Ensure that the price is a float
            price = float(pints_quarts_prices[item_data['pints_quarts_choice']][item_data['pints_quarts_mole_flavor']])
            # Use the calculated price in the f-string
            pints_quarts_ordered = (
                f"{item_data['pints_quarts_quantity']} "
                f"{item_data['pints_quarts_choice'].replace('-', ' ').replace('s', '') if item_data['pints_quarts_quantity'] == 1 else item_data['pints_quarts_choice'].replace('-', ' ')}"
                "{}... {:.2f}".format(item_data['pints_quarts_mole_flavor'], price)

            )
            pints_quarts_paragraph = doc.add_paragraph(pints_quarts_ordered, style='Normal')
            pints_quarts_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add chafers details (if any)
    if chafers_data:
        for chafer, quantity in chafers_data.items():
            # Ensure that the price is a float
            price = float(quantity * chafers_prices[chafer])
            # Determine chafer size
            chafer_size = "1/2" if "1/2" in chafer else ""
            # Determine pluralization based on quantity
            is_plural = 's' if quantity != 1 else ''
            # Use the calculated price in the f-string
            chafer_ordered = '{} {} Chafer Set{}... {:.2f}'.format(quantity, chafer_size, is_plural, price)
            chafer_paragraph = doc.add_paragraph(chafer_ordered, style='Normal')
            chafer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

             # Add a blank line
    doc.add_paragraph("", style='Normal')

         # Add a centered line with "~"
    centered_line_paragraph = doc.add_paragraph("~")
    centered_line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Calculate the Food Total by subtracting fees from combined_total
    food_total = combined_total - full_service_fee - mileage_fee - packaging_handling_fee

    # Add Food Total to the document if there are additional fees
    if mileage_fee or full_service_fee or packaging_handling_fee_percentage > 0:
        food_total_text = f"Food Total: {'%.2f' % food_total}"
        food_total_paragraph = doc.add_paragraph(food_total_text, style='Normal')
        food_total_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        # Check if Full Service is selected and display the Full Service Fee
    if 'full_service_fee' in locals() and full_service_fee > 0:
        full_service_fee_text = f"37% Full Service Fee: {'%.2f' % full_service_fee}"
        full_service_paragraph = doc.add_paragraph(full_service_fee_text, style='Normal')
        full_service_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

         # Display Packaging & Handling Fee if a percentage is provided
    if packaging_handling_fee_percentage > 0:
        packaging_handling_fee = (packaging_handling_fee_percentage / 100) * subtotal
        packaging_fee_text = f"{packaging_handling_fee_percentage:.0f}% Packaging and Handling Fee: {'%.2f' % packaging_handling_fee}"
        packaging_fee_paragraph = doc.add_paragraph(packaging_fee_text, style='Normal')
        packaging_fee_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Display Mileage Fee if provided
    if mileage_fee > 0:
        mileage_fee_text = f"Mileage Fee @.67¢: {'%.2f' % mileage_fee}"
        mileage_paragraph = doc.add_paragraph(mileage_fee_text, style='Normal')
        mileage_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Calculate and display the Subtotal
    subtotal = combined_total  # Use combined_total as the subtotal
    subtotal_text = f"Subtotal: {'%.2f' % subtotal}"
    subtotal_paragraph = doc.add_paragraph(subtotal_text, style='Normal')
    subtotal_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    # Display Tax Exempt status
    if is_tax_exempt:
        tax_exempt_paragraph = doc.add_paragraph("Tax Exempt", style='Normal')
        tax_exempt_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

   

    # Calculate and display Tax if not Tax Exempt
    if not is_tax_exempt:
        tax_rate = 0.0875
        tax_amount = (subtotal) * tax_rate
        tax_text = f"Tax {tax_rate * 100:.2f}%: {'%.2f' % tax_amount}"
        tax_paragraph = doc.add_paragraph(tax_text, style='Normal')
        tax_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Display the Total
    total_text = f"Total: {'%.2f' % total_after_tax}"
    total_paragraph = doc.add_paragraph(total_text, style='Normal')
    total_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Display Gratuity Fee if a percentage is provided
    if gratuity_fee_percentage > 0:
        gratuity_fee_text = f"{gratuity_fee_percentage:.0f}% Gratuity Fee: {gratuity_fee_formatted}"
        gratuity_fee_paragraph = doc.add_paragraph(gratuity_fee_text, style='Normal')
        gratuity_fee_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Add a centered line with "~"
    centered_line_paragraph = doc.add_paragraph("~")
    centered_line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    image_path = 'static/RI Logo.png'
        # Add the logo with custom formatting
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(image_path, width=Pt(100), height=Pt(150))

    # Add a centered line 
    centered_line_paragraph = doc.add_paragraph("~Gracias! Provecho! Eat Red Iguana!~")
    centered_line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
 

 # Add a centered line 
    centered_line_paragraph = doc.add_paragraph("~~Gratuities appreciated for driver attendant~")
    centered_line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
 

    # Generate a filename based on customer name and selected options
    filename_parts = [customer_name]

    if delivery_selection:
        filename_parts.append(delivery_selection)

    if pickup_location_selection:
        filename_parts.append(pickup_location_selection)

    if full_service_fee:
        filename_parts.append("FullService")

    if total_chafers:
        filename_parts.append("withChafers")

    # Combine filename parts with underscores and add the .docx extension
    filename = '_'.join(filename_parts) + '.docx'

    # Save the Word document to the dynamically generated filename
    doc_file = 'order_confirmation.docx'
    doc.save(doc_file)

    # Send the Word document as a downloadable file to the user
    return send_file(doc_file, as_attachment=True, download_name=filename)




if __name__ == "__main__":
    app.run()
